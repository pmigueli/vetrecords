import logging
from typing import Protocol

logger = logging.getLogger(__name__)


class TextExtractor(Protocol):
    """Interface for text extraction from different file formats."""

    def extract(self, file_path: str) -> str: ...


class PDFExtractor:
    """Extract text from PDF files using PyMuPDF."""

    def extract(self, file_path: str) -> str:
        import fitz  # PyMuPDF

        text_parts = []
        with fitz.open(file_path) as doc:
            for page_num, page in enumerate(doc):
                page_text = page.get_text()
                if page_text.strip():
                    text_parts.append(page_text)
                logger.debug(f"PDF page {page_num + 1}: {len(page_text)} chars")

        full_text = "\n".join(text_parts)
        logger.info(f"PDF extracted: {len(full_text)} chars from {len(text_parts)} pages")
        return full_text


class DOCXExtractor:
    """Extract text from DOCX files using python-docx."""

    def extract(self, file_path: str) -> str:
        from docx import Document

        doc = Document(file_path)
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    text_parts.append(row_text)

        full_text = "\n".join(text_parts)
        logger.info(f"DOCX extracted: {len(full_text)} chars")
        return full_text


class ImageExtractor:
    """Extract text from images using Tesseract OCR."""

    def extract(self, file_path: str, language_hint: str = "spa+eng") -> str:
        import pytesseract
        from PIL import Image

        image = Image.open(file_path)
        text = pytesseract.image_to_string(image, lang=language_hint)
        logger.info(f"Image OCR extracted: {len(text)} chars")
        return text


class UnsupportedFormatError(Exception):
    pass


def get_extractor(content_type: str) -> TextExtractor:
    """Factory: returns the right extractor based on file MIME type."""
    extractors: dict[str, TextExtractor] = {
        "application/pdf": PDFExtractor(),
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": DOCXExtractor(),
        "image/jpeg": ImageExtractor(),
        "image/png": ImageExtractor(),
    }
    extractor = extractors.get(content_type)
    if not extractor:
        raise UnsupportedFormatError(f"No extractor for {content_type}")
    return extractor
